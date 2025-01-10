import json
import os
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm, Pt

# Load configuration
with open('config.json', 'r') as config_file:
    config = json.load(config_file)

# Replace placeholder with actual path
input_file = (config['content_path'])
output_path = os.path.expandvars(config['output_path'])
image_path = os.path.expandvars(config['image_path'])

# Parse the XML file
tree = ET.parse(input_file)
root = tree.getroot()

# Create a new Document
doc = Document()

# Create a table with one row and two columns
table = doc.add_table(rows=1, cols=2)
table.autofit = False

# Set the width of the first cell to make more space for the image
table.columns[0].width = Cm(12)
table.columns[1].width = Cm(4)

# Add the heading and first four paragraphs to the first cell
cell = table.cell(0, 0)
cell.text = root.find('heading').text + '\n\n'
paragraphs = root.findall('paragraph')[:4]
for p in paragraphs:
    cell.text += p.text + '\n'

# Add the image to the second cell and align it to the right
cell = table.cell(0, 1)
paragraph = cell.paragraphs[0]
run = paragraph.add_run()
run.add_picture(image_path, width=Cm(4))  # Set the width of the image

# Add padding to the image to move it further to the right
paragraph_format = paragraph.paragraph_format
paragraph_format.right_indent = Cm(1)

# Iterate through the remaining XML elements and add content to the Document
for element in root:
    if element.tag == 'heading' and element.attrib['level'] != '1':
        doc.add_heading(element.text, level=int(element.attrib['level']))
    elif element.tag == 'paragraph' and element not in paragraphs:
        # Replace \n with actual newlines
        text = element.text.replace('\\n', '\n')
        doc.add_paragraph(text)

# Save the document
doc.save(output_path)